"""
Tests for Step 4 — DOCX + URL Ingestion.

Validates:
- DOCX extractor returns non-empty text from a minimal .docx fixture.
- URL endpoint rejects non-http/https schemes (e.g. file://, ftp://).
- URL endpoint rejects localhost and private IP ranges (SSRF guard).
- URL endpoint accepts a valid public URL (mocked).
"""

import io
import ipaddress
from unittest.mock import AsyncMock, MagicMock, patch

import pytest


# ── Helpers ───────────────────────────────────────────────────────────────────

def _make_minimal_docx() -> bytes:
    """Create a minimal in-memory DOCX with one paragraph using python-docx."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is a test paragraph in a DOCX file.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── DOCX extractor tests ──────────────────────────────────────────────────────

class TestDocxExtractor:
    """Tests for get_docx_pages()."""

    def test_extracts_text_from_minimal_docx(self, tmp_path):
        """get_docx_pages returns non-empty text from a minimal .docx fixture."""
        docx_path = tmp_path / "test.docx"
        docx_path.write_bytes(_make_minimal_docx())

        from backend.services.ingest import get_docx_pages

        result = get_docx_pages(str(docx_path))

        assert len(result) == 1, "Expected exactly one page"
        text, page_num = result[0]
        assert "test paragraph" in text.lower()
        assert page_num == 1

    def test_returns_empty_for_nonexistent_file(self):
        """get_docx_pages returns an empty list for a missing file."""
        from backend.services.ingest import get_docx_pages

        result = get_docx_pages("/tmp/does_not_exist_xyz.docx")
        assert result == []

    def test_extracts_table_rows(self, tmp_path):
        """get_docx_pages includes table cell content."""
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"
        buf = io.BytesIO()
        doc.save(buf)
        docx_path = tmp_path / "table.docx"
        docx_path.write_bytes(buf.getvalue())

        from backend.services.ingest import get_docx_pages

        result = get_docx_pages(str(docx_path))
        assert len(result) == 1
        text, _ = result[0]
        assert "Header A" in text
        assert "Value 2" in text


# ── URL SSRF guard tests ──────────────────────────────────────────────────────

class TestUrlSsrfGuard:
    """Tests for _check_url_safe()."""

    def test_accepts_valid_https_url(self):
        """A standard public HTTPS URL is accepted."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("https://example.com/document")
        assert safe is True
        assert reason == ""

    def test_accepts_valid_http_url(self):
        """A standard public HTTP URL is accepted."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("http://example.com/page")
        assert safe is True

    def test_rejects_file_scheme(self):
        """file:// URLs are rejected."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("file:///etc/passwd")
        assert safe is False
        assert "scheme" in reason.lower() or "file" in reason.lower()

    def test_rejects_ftp_scheme(self):
        """ftp:// URLs are rejected."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("ftp://ftp.example.com/file")
        assert safe is False

    def test_rejects_localhost(self):
        """localhost is blocked."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("http://localhost:8080/admin")
        assert safe is False

    def test_rejects_loopback_ip(self):
        """127.x.x.x loopback IPs are blocked."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("http://127.0.0.1/internal")
        assert safe is False

    def test_rejects_private_10_range(self):
        """10.0.0.0/8 private IP range is blocked."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("http://10.0.0.1/secret")
        assert safe is False

    def test_rejects_private_192_168_range(self):
        """192.168.x.x private IP range is blocked."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("http://192.168.1.100/data")
        assert safe is False

    def test_rejects_link_local_aws_metadata(self):
        """169.254.169.254 (AWS metadata endpoint) is blocked."""
        from backend.routers.knowledge import _check_url_safe

        safe, reason = _check_url_safe("http://169.254.169.254/latest/meta-data/")
        assert safe is False


# ── URL endpoint integration tests ────────────────────────────────────────────

class TestIngestUrlEndpoint:
    """Basic API-level tests for POST /ingest/url."""

    def test_rejects_non_http_scheme(self, client):
        """The endpoint returns 400 for non-http/https schemes."""
        resp = client.post("/ingest/url", json={"url": "file:///etc/passwd"})
        assert resp.status_code == 400

    def test_rejects_localhost_url(self, client):
        """The endpoint returns 400 for localhost URLs."""
        resp = client.post("/ingest/url", json={"url": "http://localhost/admin"})
        assert resp.status_code == 400

    def test_rejects_private_ip(self, client):
        """The endpoint returns 400 for private IP addresses."""
        resp = client.post("/ingest/url", json={"url": "http://192.168.1.1/data"})
        assert resp.status_code == 400

    def test_rejects_empty_url(self, client):
        """The endpoint returns 400 for an empty URL."""
        resp = client.post("/ingest/url", json={"url": ""})
        assert resp.status_code == 400
