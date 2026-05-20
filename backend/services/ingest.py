"""
Durable document ingestion pipeline using DBOS.

Each heavy step (file listing, PDF extraction, embedding, saving) is
decorated as a ``@DBOS.step()`` so that the workflow can resume from
the last completed step after a crash.
"""

import csv
import hashlib
import json
import os
import re
from typing import Dict, List, Tuple

import faiss
import numpy as np
import ollama
from dbos import DBOS
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
from pypdf import PdfReader

from backend.config import get_settings, logger
from backend.services.vector_db import vector_db

settings = get_settings()

# ── DBOS instance ────────────────────────────────────────────────────────────
# Uses a programmatic config dict derived from our centralized Settings.
dbos = DBOS(config={
    "name": "cognivault",
    "system_database_url": settings.db_url,
    "telemetry": {"logs": {"level": "info"}},
})

# Only auto-launch when running this module directly (for standalone ingestion).
if __name__ == "__main__":
    dbos.launch()


# ── Helpers ──────────────────────────────────────────────────────────────────

def _file_sha256(path: str) -> str:
    """Return the SHA-256 hex digest of a file's contents."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65_536), b""):
            h.update(block)
    return h.hexdigest()


def get_pdf_pages(path: str) -> List[Tuple[str, int]]:
    """Extract text from a PDF file, keeping track of page numbers."""
    try:
        reader = PdfReader(path)
        pages: List[Tuple[str, int]] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text = re.sub(r"\s+", " ", text).strip()
                if text:
                    pages.append((text, i + 1))
        return pages
    except Exception:
        logger.exception("Error reading PDF %s", path)
        return []

def get_text_pages(path: str) -> List[Tuple[str, int]]:
    """Extract text from a plain text, markdown, or csv file."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
            if text.strip():
                return [(text.strip(), 1)]
    except Exception:
        logger.exception("Error reading text file %s", path)
    return []


_MD_HEADERS = [("#", "H1"), ("##", "H2"), ("###", "H3")]


def get_markdown_chunks(path: str) -> List[Tuple[str, int]]:
    """Split a Markdown file on its header structure.

    Each chunk corresponds to one section bounded by H1/H2/H3 headers.
    A breadcrumb prefix ``[Section: H1 > H2 > H3]`` is prepended so the
    embedding always contains the section hierarchy even if the raw content
    gives no context.  Sections longer than ``settings.chunk_size`` will be
    split again by the generic splitter in ``ingest_workflow()``.

    Falls back to ``get_text_pages()`` on parse errors or empty results.
    """
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        if not text.strip():
            return []

        splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=_MD_HEADERS,
            strip_headers=True,   # headers go into doc.metadata; content stays clean
        )
        docs = splitter.split_text(text)

        chunks: List[Tuple[str, int]] = []
        for i, doc in enumerate(docs):
            meta = doc.metadata            # e.g. {"H1": "Intro", "H2": "Overview"}
            content = doc.page_content.strip()
            if not content:
                continue

            # Build the full breadcrumb from outer → inner heading.
            breadcrumb = " > ".join(
                meta[level] for level in ("H1", "H2", "H3") if level in meta
            )
            chunk_text = f"[Section: {breadcrumb}]\n\n{content}" if breadcrumb else content
            chunks.append((chunk_text, i + 1))

        return chunks if chunks else get_text_pages(path)
    except Exception:
        logger.exception("Error parsing Markdown %s", path)
        return get_text_pages(path)


# Number of CSV data rows bundled into each chunk.  Keep small enough that a
# chunk (header + N rows) usually fits within settings.chunk_size characters.
_CSV_ROWS_PER_CHUNK = 20


def get_csv_chunks(path: str) -> List[Tuple[str, int]]:
    """Split a CSV file into header-prefixed row-group chunks.

    Every chunk starts with the header row so the LLM always knows the column
    names, even when the chunk is not the first one.  Falls back to the plain
    text extractor if the file cannot be parsed as CSV.
    """
    try:
        with open(path, "r", encoding="utf-8", newline="") as f:
            rows = list(csv.reader(f))
    except Exception:
        logger.exception("Error reading CSV %s", path)
        return get_text_pages(path)

    if not rows:
        return []

    header = rows[0]
    data_rows = rows[1:]
    header_line = ", ".join(str(c) for c in header)

    if not data_rows:
        return [(header_line, 1)] if header_line.strip() else []

    chunks: List[Tuple[str, int]] = []
    for chunk_num, start in enumerate(range(0, len(data_rows), _CSV_ROWS_PER_CHUNK), 1):
        batch = data_rows[start : start + _CSV_ROWS_PER_CHUNK]
        lines = [header_line] + [", ".join(str(c) for c in row) for row in batch]
        chunk_text = "\n".join(lines).strip()
        if chunk_text:
            chunks.append((chunk_text, chunk_num))

    return chunks if chunks else get_text_pages(path)


def get_docx_pages(path: str) -> List[Tuple[str, int]]:
    """Extract text from a DOCX file (paragraphs + table cells) as a single page."""
    try:
        from docx import Document  # python-docx
        doc = Document(path)
        parts: List[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Table cells (each row as a pipe-delimited line)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        full_text = "\n".join(parts)
        if full_text.strip():
            return [(full_text.strip(), 1)]
        return []
    except Exception:
        logger.exception("Error reading DOCX %s", path)
        return []


# ── Durable workflow steps ───────────────────────────────────────────────────

_ALLOWED_EXTENSIONS = (".pdf", ".txt", ".md", ".csv", ".docx")


@DBOS.step()
def list_document_files() -> List[str]:
    """Return filenames that need (re-)ingestion: new files or content-changed files.

    Change detection uses a SHA-256 content hash stored per chunk.  A file is
    re-ingested when:
    - It has never been indexed (new file), or
    - Its current hash differs from the stored hash (edited file).

    When a changed file is detected its old chunks are soft-deleted
    (``deleted=True``) so stale vectors are excluded from future searches while
    the FAISS index structure remains intact.
    """
    # Build filename → stored_hash from existing metadata.
    # A filename present in metadata but without a hash (legacy chunks) is
    # treated as "unknown hash" → will be re-ingested to add the hash.
    stored_hashes: Dict[str, str | None] = {}
    if os.path.exists(settings.metadata_file):
        try:
            with open(settings.metadata_file, "r") as f:
                meta = json.load(f)
            for item in meta:
                fname = item.get("source")
                if fname and fname not in stored_hashes:
                    stored_hashes[fname] = item.get("file_hash")  # may be None
        except Exception:
            logger.exception("Failed to read existing metadata")

    files_to_ingest: List[str] = []
    if not os.path.exists(settings.docs_dir):
        return files_to_ingest

    for filename in os.listdir(settings.docs_dir):
        if not any(filename.lower().endswith(ext) for ext in _ALLOWED_EXTENSIONS):
            continue

        path = os.path.join(settings.docs_dir, filename)
        current_hash = _file_sha256(path)

        if filename not in stored_hashes:
            # Brand-new file — never seen before.
            files_to_ingest.append(filename)
        elif stored_hashes[filename] != current_hash:
            # File was edited (or previously indexed without a hash).
            # Soft-delete stale chunks so they are excluded from search.
            deleted = vector_db.delete_by_source(filename)
            logger.info(
                "File '%s' changed (old hash=%s, new hash=%s); "
                "soft-deleted %d old chunk(s)",
                filename,
                stored_hashes[filename],
                current_hash,
                deleted,
            )
            files_to_ingest.append(filename)
        # else: hash unchanged → skip

    logger.info(
        "Found %d file(s) to ingest (%d already up-to-date)",
        len(files_to_ingest),
        len(stored_hashes) - len(files_to_ingest),
    )
    return sorted(files_to_ingest)


@DBOS.step()
def process_single_document(filename: str) -> List[Dict]:
    """Process a single document and extract content.

    Each returned page dict includes a ``file_hash`` field (SHA-256 of the
    raw file bytes) so chunks stored later can be compared on re-ingest.
    """
    path = os.path.join(settings.docs_dir, filename)
    logger.info("Processing document: %s", filename)

    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        pages = get_pdf_pages(path)
        doc_type = "pdf"
        min_chunk_length = 100
    elif ext == ".docx":
        pages = get_docx_pages(path)
        doc_type = "docx"
        min_chunk_length = 100
    elif ext == ".txt":
        pages = get_text_pages(path)
        doc_type = "text"
        min_chunk_length = 100
    elif ext == ".md":
        pages = get_markdown_chunks(path)
        doc_type = "markdown"
        min_chunk_length = 20   # short sections (e.g. "## Contact\n\ninfo@…") are valid
    elif ext == ".csv":
        pages = get_csv_chunks(path)
        doc_type = "csv"
        min_chunk_length = 20   # header + a handful of rows may be brief
    else:
        logger.warning("Unsupported file type: %s", ext)
        return []

    file_hash = _file_sha256(path)

    docs: List[Dict] = []
    for text, page_num in pages:
        docs.append({
            "source": filename,
            "content": text,
            "type": doc_type,
            "page": page_num,
            "file_hash": file_hash,
            "min_chunk_length": min_chunk_length,
        })
    return docs


@DBOS.step()
def embed_batch(batch: List[str]) -> List[List[float]]:
    """Embed a batch of text chunks using the Ollama embedding model."""
    try:
        response = ollama.embed(model=settings.embedding_model, input=batch)
        return response["embeddings"]
    except Exception:
        logger.exception("Embedding batch failed — DBOS will retry")
        raise


@DBOS.step()
def save_vector_store(embeddings: List[List[float]], chunks_metadata: List[Dict]) -> None:
    """Save the FAISS index and metadata to disk (append or create)."""
    if not embeddings:
        return

    embeddings_np = np.array(embeddings).astype("float32")
    faiss.normalize_L2(embeddings_np)

    existing_meta: List[Dict] = []
    if os.path.exists(settings.index_file) and os.path.exists(settings.metadata_file):
        index = faiss.read_index(settings.index_file)
        index.add(embeddings_np)
        with open(settings.metadata_file, "r") as f:
            existing_meta = json.load(f)
    else:
        index = faiss.IndexFlatIP(embeddings_np.shape[1])
        index.add(embeddings_np)

    combined_metadata = existing_meta + chunks_metadata

    faiss.write_index(index, settings.index_file)
    with open(settings.metadata_file, "w") as f:
        json.dump(combined_metadata, f)

    logger.info("Saved %d new chunks to vector store", len(chunks_metadata))


# ── Main workflow ────────────────────────────────────────────────────────────

@DBOS.workflow()
def ingest_workflow() -> int:
    """Durable workflow: chunk → embed → store documents from docs/."""
    filenames = list_document_files()
    if not filenames:
        logger.info("No new documents found to ingest")
        return 0

    # Step 2: Extract content from each PDF
    all_documents: List[Dict] = []
    for filename in filenames:
        docs = process_single_document(filename)
        all_documents.extend(docs)

    if not all_documents:
        return 0

    # Step 3: Chunk the extracted text
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )

    chunks_metadata: List[Dict] = []
    texts_to_embed: List[str] = []

    logger.info("Chunking %d document pages", len(all_documents))
    for doc in all_documents:
        chunks = splitter.split_text(doc["content"])
        # Structured types (markdown sections, CSV row-groups) use a lower
        # minimum so short-but-meaningful chunks aren't discarded.
        min_len = doc.get("min_chunk_length", 100)
        for i, chunk in enumerate(chunks):
            if len(chunk.strip()) < min_len:
                continue
            chunks_metadata.append({
                "source": doc["source"],
                "type": doc["type"],
                "content": chunk,
                "chunk_id": i,
                "page": doc["page"],
                "file_hash": doc.get("file_hash"),   # SHA-256 for change detection
            })
            texts_to_embed.append(chunk)

    # Step 4: Embed in batches
    logger.info("Generating embeddings for %d chunks", len(texts_to_embed))
    embeddings: List[List[float]] = []
    batch_size = settings.embedding_batch_size
    for i in range(0, len(texts_to_embed), batch_size):
        batch = texts_to_embed[i : i + batch_size]
        batch_embeddings = embed_batch(batch)
        embeddings.extend(batch_embeddings)

    # Step 5: Save
    save_vector_store(embeddings, chunks_metadata)
    return len(chunks_metadata)


if __name__ == "__main__":
    count = ingest_workflow()
    logger.info("Ingestion complete — indexed %d chunks", count)
