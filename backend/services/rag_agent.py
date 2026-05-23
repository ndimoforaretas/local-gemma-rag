"""
RAG Agent — Strands SDK agent with streaming support.

Session isolation
-----------------
Each chat session gets its own isolated conversation history stored in
`_session_histories` (keyed by session_id).  A fresh Agent is created per
request but pre-loaded with the session's history, so multi-turn conversation
works correctly without sessions bleeding into each other.

Two-phase streaming
-------------------
Phase 1 — Thinking (optional):
    Direct ollama.AsyncClient call with options={"thinking": True}.
    Now correctly includes:
      - The full system prompt (so the model knows it is CogniVault)
      - Any attached images (so reasoning matches what the model sees)
    Emits  {"type": "thinking", "data": "<reasoning chunk>"}  events.

Phase 2 — Agent response:
    Normal Strands Agent stream with tool support (search_knowledge_base,
    list_documents, analyze_document, compare_documents, calculator,
    current_time).  Emits {"type": "text"} and {"type": "metadata"} events.
"""

import asyncio
import base64
import io
import json
import os
from typing import AsyncGenerator, Optional

import ollama as _ollama

from strands import Agent
from strands.models.ollama import OllamaModel

from backend.config import get_settings, logger
from backend.tools.agent_tools import (
    _last_doc_ctx,
    _source_filter_ctx,
    analyze_document,
    calculator,
    compare_documents,
    current_time,
    list_documents,
    search_knowledge_base,
)

settings = get_settings()

def _make_ollama_model() -> OllamaModel:
    """Create a fresh OllamaModel for a single request.

    Constructing per-request (rather than sharing a module-level singleton)
    ensures no per-call state on the model object is shared across concurrent
    sessions.  OllamaModel holds only configuration so construction is cheap.
    """
    return OllamaModel(
        host=settings.ollama_host,
        model_id=settings.llm_model,
        # Disable Ollama-level thinking in Phase 2 (Strands agent call).
        # Phase 1 already streams thinking via a direct ollama.AsyncClient call
        # with options={"thinking": True}.  Without this flag, Gemma 4's default
        # modelfile may still emit <think>…</think> tokens inside message.content,
        # which causes responses to appear truncated (text before the closing tag
        # is swallowed by the markdown renderer on the frontend).
        options={"thinking": False},
    )

# ── Session isolation ─────────────────────────────────────────────────────────

# Per-session conversation histories.
# Keys: session_id strings from the frontend.
# Values: Strands/Bedrock-format message lists.
_session_histories: dict[str, list] = {}

# Per-session asyncio locks: prevent two concurrent requests in the same session
# from corrupting each other's history.
_session_locks: dict[str, asyncio.Lock] = {}
_locks_meta_lock = asyncio.Lock()

# Character budget for stored history.  Rough heuristic: 4 chars ≈ 1 token.
# 24 000 chars ≈ 6 000 tokens — leaves the majority of the 128K context for
# the current query, retrieved chunks, and generation.
_MAX_HISTORY_CHARS = 24_000

# ── Attachment helpers ────────────────────────────────────────────────────────

_TEXT_MIME_TYPES = {
    "application/json",
    "application/xml",
    "application/yaml",
    "application/x-yaml",
    "application/javascript",
    "application/typescript",
    "application/csv",
}
_TEXT_FILE_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".csv", ".json", ".xml",
    ".yaml", ".yml", ".log", ".py", ".js", ".ts", ".tsx", ".jsx", ".sql",
}

_PDF_MIME = "application/pdf"
_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/docx",
}

# Maximum characters extracted from a single PDF/DOCX in chat context.
# Keeps the prompt manageable; the full document can be indexed in the KB.
_MAX_EXTRACTED_CHARS = 15_000


async def _get_session_lock(session_id: str) -> asyncio.Lock:
    """Return (creating if needed) the per-session asyncio Lock."""
    async with _locks_meta_lock:
        if session_id not in _session_locks:
            _session_locks[session_id] = asyncio.Lock()
        return _session_locks[session_id]


def _trim_history(history: list) -> list:
    """
    Drop the oldest user/assistant turn-pairs until the total character count
    of the history fits within _MAX_HISTORY_CHARS.
    """
    def _char_count(h: list) -> int:
        total = 0
        for msg in h:
            content = msg.get("content", "")
            if isinstance(content, str):
                total += len(content)
            elif isinstance(content, list):
                for block in content:
                    if isinstance(block, dict):
                        total += len(str(block.get("text", "")))
        return total

    # Drop pairs (user + assistant) from the front until we fit.
    while _char_count(history) > _MAX_HISTORY_CHARS and len(history) >= 2:
        history = history[2:]
    return history


async def _stream_thinking(
    text_query: str,
    image_bytes_list: Optional[list[bytes]] = None,
    system_prompt: Optional[str] = None,
) -> AsyncGenerator[str, None]:
    """
    Phase 1: stream Gemma 4's internal reasoning chain.

    Sends the full system prompt plus the user's text and any attached images
    so the reasoning panel reflects exactly what the model will actually see.

    Emits {"type": "thinking", "data": "<chunk>"} JSON Lines events.
    Yields nothing if thinking_mode is disabled or the model returns no tokens.
    """
    if not settings.thinking_mode:
        return

    client = _ollama.AsyncClient(host=settings.ollama_host)
    try:
        user_msg: dict = {"role": "user", "content": text_query}
        if image_bytes_list:
            user_msg["images"] = image_bytes_list

        stream = await client.chat(
            model=settings.llm_model,
            messages=[
                {"role": "system", "content": system_prompt or settings.agent_system_prompt},
                user_msg,
            ],
            options={"thinking": True},
            stream=True,
        )
        async for chunk in stream:
            thinking_chunk = chunk.message.thinking
            if thinking_chunk:
                yield f'{json.dumps({"type": "thinking", "data": thinking_chunk})}\n'
    except Exception:
        # Thinking is best-effort — never block the main Phase 2 response.
        logger.warning("Thinking phase skipped (model or connection error).")
        return


def _mime_to_format(mime_type: str) -> str:
    """Convert 'image/png' → 'png' for Strands image content blocks."""
    mapping = {
        "image/jpeg": "jpeg",
        "image/jpg": "jpeg",
        "image/png": "png",
        "image/gif": "gif",
        "image/webp": "webp",
    }
    return mapping.get(mime_type, mime_type.split("/")[-1])


def _is_text_attachment(att: object) -> bool:
    """Treat common text-like files (including PDF/DOCX) as text attachments."""
    mime_type = (getattr(att, "mime_type", "") or "").lower()
    if mime_type.startswith("text/") or mime_type in _TEXT_MIME_TYPES:
        return True
    if mime_type == _PDF_MIME or mime_type in _DOCX_MIME_TYPES:
        return True
    name = getattr(att, "name", None) or ""
    _, ext = os.path.splitext(name.lower())
    return ext in _TEXT_FILE_EXTENSIONS or ext in (".pdf", ".docx")


def _extract_text_from_attachment(att: object, raw_bytes: bytes) -> str:
    """
    Extract readable text from any supported attachment type.

    Handles: plain text, PDF (via pypdf), DOCX (via python-docx).
    Returns the extracted text, truncated to _MAX_EXTRACTED_CHARS.
    """
    mime_type = (getattr(att, "mime_type", "") or "").lower()
    name = getattr(att, "name", None) or ""
    _, ext = os.path.splitext(name.lower())

    # ── PDF ──────────────────────────────────────────────────────────────────
    if mime_type == _PDF_MIME or ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            pages = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(page_text)
            text = "\n\n".join(pages)
        except Exception as exc:
            logger.warning("PDF extraction failed for '%s': %s", name, exc)
            text = "[PDF content could not be extracted]"
        return text[:_MAX_EXTRACTED_CHARS]

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if mime_type in _DOCX_MIME_TYPES or ext == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(raw_bytes))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            text = "\n".join(parts)
        except Exception as exc:
            logger.warning("DOCX extraction failed for '%s': %s", name, exc)
            text = "[DOCX content could not be extracted]"
        return text[:_MAX_EXTRACTED_CHARS]

    # ── Plain text / source code ──────────────────────────────────────────────
    try:
        text = raw_bytes.decode("utf-8", errors="replace")
    except Exception:
        text = "[File content could not be decoded]"
    return text[:_MAX_EXTRACTED_CHARS]


async def run_rag_stream(
    query: str,
    attachments: Optional[list] = None,
    session_id: Optional[str] = None,
    document_filter: Optional[list[str]] = None,
    trim_history_to_turns: Optional[int] = None,
) -> AsyncGenerator[str, None]:
    """
    Stream agentic RAG responses using JSON Lines format.

    Yields one JSON object per line:
    - {"type": "thinking",  "data": "<reasoning chunk>"}   (Phase 1, optional)
    - {"type": "text",      "data": "text chunk"}           (Phase 2)
    - {"type": "metadata",  "data": {source document meta}} (Phase 2)
    - {"type": "error",     "data": "message"}              (on failure)

    Parameters
    ----------
    session_id
        Frontend-assigned session identifier used to isolate conversation
        history between chat sessions.  When None, a shared anonymous bucket
        is used (safe for single-request / test scenarios).
    """
    _last_doc_ctx.set([])
    # Apply document-scope filter for this request (ContextVar is Task-local).
    _source_filter_ctx.set(document_filter if document_filter else None)

    # Build a dynamic system prompt that tells the agent about the active scope.
    try:
        if document_filter:
            cats: set[str] = set()
            filter_set = set(document_filter)
            for chunk in vector_db.metadata:
                if not chunk.get("deleted") and chunk.get("source") in filter_set:
                    cat = chunk.get("category", "General")
                    if cat:
                        cats.add(cat)
            scope_desc = (
                f"category '{', '.join(sorted(cats))}'"
                if cats
                else f"document(s): {', '.join(list(document_filter)[:5])}"
            )
            scope_note = (
                f"\n\nSEARCH SCOPE ACTIVE — MANDATORY OVERRIDE: The user has pinned the search "
                f"to {scope_desc}. This overrides all other search rules, including the 'skip "
                f"search for general questions' rule. You MUST call search_knowledge_base FIRST "
                f"for every question when a scope is active — no exceptions, even for questions "
                f"you think you know the answer to. "
                f"search_knowledge_base will ONLY return results from the selected document(s). "
                f"Begin your response with 'Based on [{scope_desc}]:'. "
                f"If the scoped document(s) do not contain enough information, state: "
                f"'The selected document(s) do not contain sufficient information on this topic.'"
            )
            effective_system_prompt = settings.agent_system_prompt + scope_note
        else:
            effective_system_prompt = settings.agent_system_prompt
    except Exception:
        logger.exception("Failed to build scope-aware system prompt; falling back to default")
        effective_system_prompt = settings.agent_system_prompt

    # ── Build prompt ──────────────────────────────────────────────────────────
    # When a scope filter is active, prepend an explicit scope header to the
    # user query.  This is far more reliable than only mentioning it in the
    # system prompt because (1) it appears in the user message the model is
    # reasoning about, (2) both Phase 1 thinking and Phase 2 agent see it,
    # (3) the model can't "decide" to ignore it because it's part of the
    # question itself.
    effective_query: str = query or ""
    if document_filter:
        try:
            cat_list = sorted(cats) if "cats" in locals() and cats else []
        except Exception:
            cat_list = []
        if cat_list:
            files_preview = ", ".join(list(document_filter)[:5])
            scope_header = (
                f"[SEARCH SCOPE: category '{', '.join(cat_list)}' — "
                f"{len(document_filter)} file(s): {files_preview}]"
            )
        else:
            scope_header = (
                f"[SEARCH SCOPE: {len(document_filter)} file(s): "
                f"{', '.join(list(document_filter)[:5])}]"
            )
        effective_query = (
            f"{scope_header}\n"
            f"You MUST call search_knowledge_base to look inside the scoped "
            f"file(s) above before answering. Do not skip the search.\n\n"
            f"Question: {query}"
        )

    user_input: object = effective_query
    combined_text: str = effective_query
    raw_image_bytes: list[bytes] = []       # collected for Phase 1 thinking call

    if attachments:
        image_blocks = []
        text_files: list[tuple[str, str]] = []

        for att in attachments:
            if att.mime_type.startswith("image/"):
                img_bytes = base64.b64decode(att.data)
                raw_image_bytes.append(img_bytes)
                image_blocks.append({
                    "image": {
                        "format": _mime_to_format(att.mime_type),
                        "source": {"bytes": img_bytes},
                    }
                })
            elif _is_text_attachment(att):
                try:
                    raw_bytes = base64.b64decode(att.data)
                    file_text = _extract_text_from_attachment(att, raw_bytes)
                    file_label = att.name or "attached file"
                    text_files.append((file_label, file_text))
                except Exception:
                    logger.warning("Failed to decode attachment '%s'", getattr(att, "name", "?"))

        if text_files:
            parts = [effective_query, ""]
            if image_blocks:
                parts.append(
                    f"[{len(image_blocks)} image(s) + {len(text_files)} file(s) attached — "
                    f"analyze ALL inputs: describe each image visually AND read each file]"
                )
            else:
                parts.append(f"[{len(text_files)} attached file(s) — read carefully and answer]")
            parts.append("")
            for i, (label, content) in enumerate(text_files, 1):
                parts.append(f"=== FILE {i}/{len(text_files)}: {label} ===")
                parts.append(content)
                parts.append(f"=== END FILE {i} ===")
                parts.append("")
            combined_text = "\n".join(parts)
        else:
            combined_text = effective_query

        if image_blocks:
            # Images first (Gemma 4 best practice: modalities before text)
            user_input = image_blocks + [{"text": combined_text}]
        elif text_files:
            user_input = combined_text

    # ── Phase 1: Thinking ─────────────────────────────────────────────────────
    # Only run if we have something non-trivial to reason about.
    if combined_text.strip() or raw_image_bytes:
        async for thinking_event in _stream_thinking(
            combined_text, raw_image_bytes or None, effective_system_prompt
        ):
            yield thinking_event

    # ── Phase 2: Agent response ───────────────────────────────────────────────
    sid = session_id or "__anonymous__"
    lock = await _get_session_lock(sid)

    async with lock:
        # Rewind history when the user edits a message or regenerates a response.
        if trim_history_to_turns is not None:
            stored = _session_histories.get(sid)
            if stored is not None:
                keep = trim_history_to_turns * 2   # each turn = 1 user + 1 assistant
                del stored[keep:]

        # Load this session's conversation history.
        history = list(_session_histories.get(sid, []))

        # Create a fresh Agent for this request (no shared mutable state).
        agent = Agent(
            model=_make_ollama_model(),
            tools=[
                search_knowledge_base,
                list_documents,
                analyze_document,
                compare_documents,
                calculator,
                current_time,
            ],
            system_prompt=effective_system_prompt,
        )

        # Restore session history into the new agent.
        if history:
            try:
                agent_msgs = getattr(agent, "messages", None)
                if agent_msgs is not None:
                    agent_msgs.clear()
                    agent_msgs.extend(history)
                else:
                    agent.messages = list(history)
            except Exception:
                pass  # Never crash on history restore; just start fresh.

        try:
            # Tracks how many citation docs have already been sent to the
            # frontend.  Tools run *before* text deltas arrive, so by the
            # time we see the first text chunk all citations are collected.
            emitted_docs = 0

            async for event in agent.stream_async(user_input):
                ev = event.get("event", {})

                # Text delta from the model response.
                delta_text = (
                    ev.get("contentBlockDelta", {}).get("delta", {}).get("text")
                )
                if delta_text:
                    # Flush any citations accumulated by search_knowledge_base
                    # since the last text chunk.  Emitting here (rather than at
                    # tool-call start) guarantees the tool has already run and
                    # _last_doc_ctx is fully populated.
                    all_docs = _last_doc_ctx.get() or []
                    while emitted_docs < len(all_docs):
                        yield f'{json.dumps({"type": "metadata", "data": all_docs[emitted_docs]})}\n'
                        emitted_docs += 1
                    yield f'{json.dumps({"type": "text", "data": delta_text})}\n'

        except Exception:
            logger.exception("Error in RAG stream for query: %s", query[:200])
            yield (
                f'{json.dumps({"type": "error", "data": "An internal error occurred while processing your query."})}\n'
            )
        finally:
            # Persist trimmed history for next request in this session.
            try:
                updated = list(getattr(agent, "messages", []))
                _session_histories[sid] = _trim_history(updated)
            except Exception:
                pass  # Never crash on history save.
